from msword_properties_generator.utils.utils_config import config  # importing centralized config
from msword_properties_generator.utils.utils_mail import safe_get
from msword_properties_generator.utils.utils_image import get_image_and_decrypt_from_image_folder
from lxml import etree as ET
from docx import Document
import tempfile
import logging
import zipfile
import shutil
import os


# Update the custom properties in de docx document structure
# ==========================================================
def set_custom_properties(extracted_dir, provider_replacements, customer_replacements):
    property_names = load_custom_property_names_map()
    # Iterate once for provider (there should only 1 provider!)
    for key_prov, value_prov in provider_replacements.items():
        if key_prov != "LeverancierEmail":
            set_custom_property(extracted_dir, property_names[key_prov], value_prov)

    for key_cust, value_cust in customer_replacements.items():
        if (key_cust != "LeverancierEmail") and (key_cust != "UploadDropbox"):
            set_custom_property(extracted_dir, property_names[key_cust], value_cust)

def set_custom_property(extracted_dir, property_name, property_value):
    custom_props_path = os.path.join(extracted_dir, 'docProps', 'custom.xml')

    tree = ET.parse(custom_props_path)
    root = tree.getroot()

    namespaces = {
        'cp': config["namespaces"]["cp"],
        'vt': config["namespaces"]["vt"]
    }

    prop = None
    for p in root.findall('cp:property', namespaces):
        if p.get('name') == property_name:
            prop = p
            break

    if prop is not None:
        vt_elem = next(iter(prop))
        vt_elem.text = str(property_value)
    else:
        prop_ids = [int(p.get('pid')) for p in root.findall('cp:property', namespaces)]
        new_pid = max(prop_ids, default=1) + 1

        prop = ET.SubElement(root, '{%s}property' % namespaces['cp'], name=property_name,
                             pid=str(new_pid), fmtid='{D5CDD505-2E9C-101B-9397-08002B2CF9AE}')
        vt_elem = ET.SubElement(prop, '{%s}lpwstr' % namespaces['vt'])
        vt_elem.text = str(property_value)

    tree.write(custom_props_path, encoding='UTF-8', xml_declaration=True, standalone=True)
    logging.debug(f"📝📋🔤 '{property_name}' successfully replaced by 'value omitted for privacy' in properties docx structure")

def extract_docx(docx_path):
    extracted_dir = tempfile.mkdtemp()
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        docx_zip.extractall(extracted_dir)
        logging.debug(f"🗜️➡️ docx successfully extracted to '{extracted_dir}' for custom properties replacements docx structure")
    return extracted_dir


def repack_docx(extracted_dir, base_document):
    save_as_docx = base_document + ".docx"
    output_docx_path = os.path.abspath(save_as_docx)
    with zipfile.ZipFile(output_docx_path, 'w', zipfile.ZIP_DEFLATED) as docx_zip:
        for foldername, subfolders, filenames in os.walk(extracted_dir):
            for filename in filenames:
                abs_name = os.path.join(foldername, filename)
                arc_name = os.path.relpath(abs_name, extracted_dir)
                docx_zip.write(abs_name, arc_name)
    shutil.rmtree(extracted_dir)

    logging.debug(f"🗜️⬅️ docx successfully repacked to 'file name omitted for privacy' after custom properties replacements docx structure")


def update_custom_properties_docx_structure(customer_line, provider_line):
    # In the custom properties xml structure...
    base_document_to_save = set_custom_properties_docx(customer_line, provider_line)
    logging.info(f"📝📷✅ Advanced properties successfully set in docx structure.")
    # Replace in the document itself...
    update_custom_properties_docx(base_document_to_save, customer_line, provider_line)
    logging.info(f"📝📷✅ Property fields successfully replaced in docx itself.")
    return base_document_to_save

def set_custom_properties_docx(customer_line, provider_line):
    extracted_dir = extract_docx(config["paths"]["word_template_path"])
    set_custom_properties(extracted_dir, provider_line, customer_line)
    base_document_to_save = build_base_document_to_save(provider_line, customer_line)
    repack_docx(extracted_dir, base_document_to_save)
    return base_document_to_save

def update_custom_properties_docx(base_document_to_save, customer_line, provider_line):
    document = open_document(base_document_to_save)
    replace_images(document, customer_line['LeverancierEmail'])
    replace_direct_text(document, provider_line, customer_line)
    save_document(base_document_to_save, document)


# Update the custom properties, images etc in de docx document itself
# ===================================================================
def replace_direct_text(document, provider_replacements, customer_replacements):
    def replace_in_paragraph(paragr, replacem):
        for run in paragr.runs:
            for old, new in replacem.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    logging.debug(f"📝📋🔤 '{old}' successfully replaced by 'new value omitted for privacy' in target file in paragraphs")

    # Replace in paragraphs
    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, provider_replacements)
        replace_in_paragraph(paragraph, customer_replacements)
    # Replace in tables (cells)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, provider_replacements)
                    replace_in_paragraph(paragraph, customer_replacements)

def replace_images_by_alt_text(doc, alt_text, new_image_path):
    logging.debug(f"📝📋📷 'Replacing image w alt_text '{alt_text}' by '{new_image_path}' in target file in paragraphs")

    if os.path.exists(new_image_path):
        with open(new_image_path, 'rb') as image_file:
            image_data = image_file.read()
        replaced = False
        for shape in doc.inline_shapes:
            doc_pr = shape._inline.docPr  # protected member accessed due to limitation of python-docx API
            if doc_pr.get('descr') == alt_text:
                rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                doc.part.related_parts[rId]._blob = image_data
                replaced = True
                pass
        if not replaced:
            logging.warning(f"📝📷⚠️ No image found with alt_text '{alt_text}'")
        else:
            logging.info(f"📝📷✅ Image with alt_text '{alt_text}' in target file successfully replaced")
    else:
        msg = f"📷❌ Error reading image file: {new_image_path}."
        logging.error(msg)
        raise ValueError(msg)

def build_base_document_to_save(provider_replacements, customer_replacements):
    leverancier_naam = safe_get(provider_replacements, "LeverancierNaam")
    klant_naam = safe_get(customer_replacements, "KlantNaam")
    klant_job_title = safe_get(customer_replacements, "KlantJobTitle")
    klant_job_reference = safe_get(customer_replacements, "KlantJobReference")
    base_document = f"{config['paths']['base_output_document_path']} - {leverancier_naam} - {klant_naam} - {klant_job_title} - {klant_job_reference}"

    return base_document

def replace_images(document, leverancier_email):
    decrypted_image_path = get_image_and_decrypt_from_image_folder(leverancier_email)
    if decrypted_image_path:
        replace_images_by_alt_text(document, config["alt_texts"]["left"], decrypted_image_path)
        replace_images_by_alt_text(document, config["alt_texts"]["right"], decrypted_image_path)

def open_document(base_document):
    save_as_docx = base_document + ".docx"
    return Document(save_as_docx)

def save_document(base_document, document):
    save_as_docx = base_document + ".docx"
    document.save(save_as_docx)

def load_custom_property_names_map():
    return{
        "KlantNaam": "Klant Naam",
        "KlantJobTitle": "Klant JobTitle",
        "KlantJobReference": "Klant JobReference",
        "LeverancierEmail": "Leverancier Email",
        "LeverancierNaam": "Leverancier Naam",
        "LeverancierStad": "Leverancier Stad",
        "LeverancierStraat": "Leverancier Straat",
        "LeverancierPostadres": "Leverancier Postadres",
        "LeverancierKandidaat": "Leverancier Kandidaat",
        "LeverancierOpgemaaktte": "Leverancier Opgemaakt te",
        "LeverancierHoedanigheid": "Leverancier Hoedanigheid",
        "UploadDropbox": "Upload Dropbox"
    }
