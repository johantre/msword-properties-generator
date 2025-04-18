from unittest.mock import patch
from docx.shared import Inches
from pathlib import Path
from docx import Document
from PIL import Image
import unittest
import tempfile
import zipfile
import shutil
import sys
import io
import os


# Add the src directory to the Python path
src_path = str(Path(__file__).resolve().parent.parent / 'src')
if src_path not in sys.path:
    sys.path.insert(0, src_path)

from msword_properties_generator.utils.utils_docx import (
    extract_docx, repack_docx,
    update_custom_properties_docx_structure,
    replace_direct_text,
    replace_images, open_document, save_document, build_base_document_to_save,
    load_custom_property_names_map
)

class TestUtilsDocx(unittest.TestCase):
    def setUp(self):
        # Create a temporary directory for test files
        self.test_dir = tempfile.mkdtemp()
        self.test_docx = os.path.join(self.test_dir, "test.docx")
        
        # Create a test Word document with proper XML structure
        doc = Document()
        doc.add_paragraph("Test paragraph with {LeverancierNaam}")
        doc.add_paragraph("Another paragraph with {KlantNaam}")
        doc.save(self.test_docx)
        
        # Extract the docx to add custom.xml
        with zipfile.ZipFile(self.test_docx, 'a') as docx_zip:
            # Create a basic custom.xml
            custom_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
            xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
    <property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="2" name="Leverancier Naam">
        <vt:lpwstr>Test Value</vt:lpwstr>
    </property>
</Properties>'''
            docx_zip.writestr('docProps/custom.xml', custom_xml)
        
        # Mock config
        self.config_patch = patch('msword_properties_generator.utils.utils_docx.config', {
            "paths": {
                "word_template_path": self.test_docx,
                "base_output_document_path": os.path.join(self.test_dir, "output"),
                "resource_path": "resources"
            },
            "namespaces": {
                "cp": "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties",
                "vt": "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"
            },
            "alt_texts": {
                "left": "Left Image",
                "right": "Right Image"
            }
        })
        self.mock_config = self.config_patch.start()
        
        # Sample test data
        self.provider_replacements = {
            "LeverancierNaam": "Test Provider",
            "LeverancierEmail": "provider@test.com",
            "LeverancierStad": "Test City",
            "LeverancierStraat": "Test Street",
            "LeverancierPostadres": "Test Address",
            "LeverancierKandidaat": "Test Candidate",
            "LeverancierOpgemaaktte": "Test Date",
            "LeverancierHoedanigheid": "Test Role",
            "UploadDropbox": "True"
        }
        
        self.customer_replacements = {
            "KlantNaam": "Test Customer",
            "KlantJobTitle": "Test Job",
            "KlantJobReference": "TEST-123",
            "LeverancierEmail": "provider@test.com"  # Added for image replacement
        }

    def tearDown(self):
        # Clean up temporary files
        shutil.rmtree(self.test_dir)
        self.config_patch.stop()

    def test_load_custom_property_names_map(self):
        property_map = load_custom_property_names_map()
        self.assertIsInstance(property_map, dict)
        self.assertEqual(property_map["LeverancierNaam"], "Leverancier Naam")
        self.assertEqual(property_map["KlantNaam"], "Klant Naam")

    def test_extract_docx(self):
        extracted_dir = extract_docx(self.test_docx)
        self.assertTrue(os.path.exists(extracted_dir))
        self.assertTrue(os.path.isdir(extracted_dir))
        # Clean up
        shutil.rmtree(extracted_dir)

    def test_repack_docx(self):
        # First extract
        extracted_dir = extract_docx(self.test_docx)
        output_path = os.path.join(self.test_dir, "output.docx")
        
        # Then repack
        repack_docx(extracted_dir, output_path.replace(".docx", ""))
        
        # Verify
        self.assertTrue(os.path.exists(output_path))
        self.assertTrue(zipfile.is_zipfile(output_path))

    @patch('msword_properties_generator.utils.utils_docx.get_image_and_decrypt_from_image_folder')
    def test_replace_images(self, mock_get_image):
        # Create a valid PNG image
        img = Image.new('RGB', (100, 100), color='red')
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_byte_arr = img_byte_arr.getvalue()
        
        # Save the image
        mock_image_path = os.path.join(self.test_dir, "test_image.png")
        with open(mock_image_path, 'wb') as f:
            f.write(img_byte_arr)
        
        mock_get_image.return_value = mock_image_path
        
        # Create test document with images
        doc = Document()
        run = doc.add_paragraph().add_run()
        run.add_picture(mock_image_path, width=Inches(1))
        doc.inline_shapes[-1]._inline.docPr.descr = "Left Image"  # Set alt text
        doc.save(self.test_docx)
        
        # Test image replacement
        doc = Document(self.test_docx)
        replace_images(doc, "test@example.com")
        
        # Verify mock was called
        mock_get_image.assert_called_once_with("test@example.com")

    def test_build_base_document_to_save(self):
        base_doc = build_base_document_to_save(self.provider_replacements, self.customer_replacements)
        expected = f"{self.mock_config['paths']['base_output_document_path']} - Test Provider - Test Customer - Test Job - TEST-123"
        self.assertEqual(base_doc, expected)

    def test_replace_direct_text(self):
        doc = Document()
        
        # Add paragraph with run
        p = doc.add_paragraph()
        run = p.add_run("Test ")
        run = p.add_run("{LeverancierNaam}")
        
        # Add table with run
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        p = cell.paragraphs[0]
        run = p.add_run("Test ")
        run = p.add_run("{LeverancierNaam}")
        
        # Create replacements with curly braces
        replacements = {"{LeverancierNaam}": "Test Provider"}
        replace_direct_text(doc, replacements, {})
        
        # Verify replacements
        self.assertEqual(doc.paragraphs[0].text, "Test Test Provider")
        self.assertEqual(doc.tables[0].cell(0, 0).text, "Test Test Provider")

    def test_open_and_save_document(self):
        # Test opening
        doc = open_document(self.test_docx.replace(".docx", ""))
        self.assertTrue(hasattr(doc, 'save'))  # Check if it has Document-like behavior
        
        # Test saving
        save_path = os.path.join(self.test_dir, "saved.docx")
        save_document(save_path.replace(".docx", ""), doc)
        self.assertTrue(os.path.exists(save_path))

    @patch('msword_properties_generator.utils.utils_docx.get_image_and_decrypt_from_image_folder')
    def test_update_custom_properties_docx_structure(self, mock_get_image):
        # Create a valid PNG image
        img = Image.new('RGB', (100, 100), color='red')
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_byte_arr = img_byte_arr.getvalue()
        
        # Save the image
        mock_image_path = os.path.join(self.test_dir, "test_image.png")
        with open(mock_image_path, 'wb') as f:
            f.write(img_byte_arr)
        
        mock_get_image.return_value = mock_image_path
        
        # Test the update
        base_doc = update_custom_properties_docx_structure(self.customer_replacements, self.provider_replacements)
        
        # Verify the document was created
        docx_path = f"{base_doc}.docx"
        self.assertTrue(os.path.exists(docx_path))
        
        # Clean up
        if os.path.exists(docx_path):
            os.remove(docx_path)

if __name__ == '__main__':
    unittest.main() 