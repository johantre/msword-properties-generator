from dropbox.exceptions import ApiError, AuthError
from dropbox.files import WriteMode
from email.message import EmailMessage
from jproperties import Properties
from docx2pdf import convert
from docx import Document
from lxml import etree as ET
import pandas as pd
import subprocess
import argparse
import tempfile
import dropbox
import logging
import smtplib
import zipfile
import shutil
import sys
import re
import os

from stone.backends.python_rsrc.stone_validators import Boolean

# Fetch properties
try:
    configs = Properties()
    with open('env/prod.properties', 'rb') as read_prop:
        configs.load(read_prop)
except (FileNotFoundError, Exception) as e:
    logging.error(f"❌Error reading properties file: {e}")
    raise SystemExit(e)

# Mail constructions
mail_smtp_port = configs.get("mail.smtp_port").data
mail_smtp_server = configs.get("mail.smtp_server").data
mail_sender_email = configs.get("mail.sender_email").data
# Path constructions
resource_path = configs.get("path.resource").data
output_path = configs.get("path.output").data
image_file_path = os.path.join(resource_path, configs.get("path.resource.image_signature").data)
dropbox_destination_folder = configs.get("path.dropbox.destination.folder").data
xls_offers_log = os.path.join(resource_path, configs.get("base.excel.offers.log").data)
xls_offers_log_sheetname = configs.get("base.excel.offers.log.sheetname").data
xls_offers_provider = os.path.join(resource_path, configs.get("base.excel.offers.provider").data)
xls_offers_provider_sheetname = configs.get("base.excel.offers.provider.sheetname").data
xls_offers_customer = os.path.join(resource_path, configs.get("base.excel.offers.customer").data)
xls_offers_customer_sheetname = configs.get("base.excel.offers.customer.sheetname").data
base_document_name = configs.get("base.word.template").data
base_document_namespace_cp = configs.get("base.word.namespace.cp").data
base_document_namespace_vt = configs.get("base.word.namespace.vt").data
base_document_image_alt_text_left =  configs.get("base.word.template.image_alt_text_left").data
base_document_image_alt_text_right =  configs.get("base.word.template.image_alt_text_right").data

word_template_path =  os.path.join(resource_path, base_document_name + '.docx')
base_output_document_path = os.path.join(output_path, base_document_name)

# Remove any existing handlers explicitly:
handlers = logging.root.handlers[:]
for handler in handlers:
    logging.root.removeHandler(handler)

logging.basicConfig(
    level=os.getenv('LOG_LEVEL', 'INFO').upper(),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)]
)


def _main(verbose=False, optional_args=None):
    logging.getLogger().handlers[0].flush()  # explicitly flush the output clearly

    def merge_replacements(customer_replacements, provider_replacements):
        combined_replacements = {}
        for key, value in provider_replacements.items():
            combined_replacements[f'prov_{key}'] = value
        for key, value in customer_replacements.items():
            combined_replacements[f'cust_{key}'] = value
        return combined_replacements

    # Read the Excel files into a DataFrames
    # log_data_frame = safely_read_excel(xls_offers_log, xls_offers_log_sheetname, "offers log")
    # provider_data_frame = safely_read_excel(xls_offers_provider, xls_offers_provider_sheetname, "offers provider")
    # customers_data_frame = safely_read_excel(xls_offers_customer, xls_offers_customer_sheetname, "offers customer")

    provider_replacements = create_sanitized_replacements(xls_offers_provider, xls_offers_provider_sheetname, "prov")
    customer_replacements = create_sanitized_replacements(xls_offers_customer, xls_offers_customer_sheetname, "cust", optional_args)
    combined_replacements = merge_replacements(customer_replacements, provider_replacements)

    # Each row in offersCustomer = new offer docx + pdf
    provider_line = None
    base_document_to_save = None
    for index, (key, line) in enumerate(combined_replacements.items()):
        if index == 0:
            # Special handling for the first item
            logging.debug(f"Running over combined_replacements. First item: {key} -> {line}")
            if isinstance(line, dict):
                provider_line = line
                base_document_to_save = ""
                for key_prov, value_prov in provider_line.items():
                    logging.debug(f'Provider: {key_prov}: {value_prov}')
            else:
                logging.warning(f'No provider dict found in line: {line[index]}')
        else:
            # Handling for all other items
            logging.debug(f"Running over combined_replacements. Item {index}: {key} -> {line}")
            if isinstance(line, dict):
                customer_line = line
                # For each customer row ...
                # In the custom properties xml structure...
                extracted_dir = extract_docx(word_template_path)
                set_custom_properties(extracted_dir, provider_line, customer_line)
                base_document_to_save = build_base_document_to_save(provider_line, customer_line)
                repack_docx(extracted_dir, base_document_to_save)

                # Replace in the document itself...
                document = open_document(base_document_to_save)
                replace_images(document)
                replace_direct_text(document, provider_line, customer_line)
                save_document(base_document_to_save, document)

                convert_to_pdf(base_document_to_save)

                # operation succeeded, append row to 'log' + drop row from 'new'
                # log_data_frame = pd.concat([log_data_frame, customer_line[index].to_frame().T], axis='index', ignore_index=True)
                # customers_data_frame.drop(labels=[index], axis='index', inplace=True)
                # Update log sheet w fresh offer
                # save_to_excel(customers_data_frame, log_data_frame)

                for key_cust, value_cust in customer_line.items():
                    logging.debug(f'Customer(s): {key_cust}: {value_cust}')
            else:
                logging.warning(f'No customer dict found in line: {line[index]}')

        if base_document_to_save.strip():
            if optional_args:
                recipient_email = optional_args["EmailRecipient"]
            else:
                recipient_email = 'johan_tre@hotmail.com'
            generated_files = [base_document_to_save + ".docx", base_document_to_save + ".pdf"]

            send_email(generated_files, recipient_email, provider_line, customer_line)

            upload_dropbox = True
            if optional_args:
                if not optional_args["UploadDropbox"] == 'true':
                    upload_dropbox = False
                    logging.info(f"ℹ️Not uploaded to Dropbox, as requested by user. Value option_args['UploadDropbox'] is: {optional_args["UploadDropbox"]}.")

            if upload_dropbox:
                try:
                    dropbox_upload(generated_files)
                except AuthError as ae:
                    logging.error(f"Authentication error while uploading dropbox: {ae}")
                    # Handle the authentication error, such as prompting for re-authentication
                except Exception as ee:
                    # Handle other possible exceptions
                    logging.error(f"An error occurred while uploading dropbox: {ee}")
            else:
                logging.info(f"ℹ️Not uploaded to Dropbox Johan!")



# Dropbox API setup:
def get_dbx_client():
    APP_KEY = os.environ.get('DROPBOX_APP_KEY')
    APP_SECRET = os.environ.get('DROPBOX_APP_SECRET')
    REFRESH_TOKEN = os.environ.get('DROPBOX_REFRESH_TOKEN')

    return dropbox.Dropbox(
        oauth2_refresh_token=REFRESH_TOKEN,
        app_key=APP_KEY,
        app_secret=APP_SECRET
    )

def dropbox_upload(generated_files):
    dbx = get_dbx_client()

    for filepath in generated_files:
        abs_full_path = os.path.abspath(filepath)
        if os.path.exists(filepath):
            logging.info(f"ℹ️File at location '{filepath}' found, at absolute path {abs_full_path}")
        else:
            logging.warning(f"⚠️File at location '{filepath}' not found, at absolute path {abs_full_path}!")

        dropbox_dest_path = os.path.join(dropbox_destination_folder, os.path.basename(filepath)).replace('\\', '/')

        logging.info(f"📂 Local file to upload: {abs_full_path}")
        logging.info(f"📌 Dropbox full destination path: {dropbox_dest_path}")
        with open(abs_full_path, 'rb') as file:
            try:
                logging.info(f"📤 Uploading '{filepath}' to '{dropbox_dest_path}' on Dropbox.")
                response = dbx.files_upload(
                    file.read(),
                    dropbox_dest_path,
                    mode=WriteMode("overwrite")
                )
                logging.info(f"🚀 Successfully uploaded file. Dropbox file details: {response}")
                print(f"✅ Successfully uploaded file to Dropbox: '{os.path.basename(filepath)}'")
            except dropbox.exceptions.ApiError as err:logging.error(f"📛 Dropbox API error: {err}")


def send_email(generated_files, email_address, provider_replacements, customer_replacements):
    # Email setup (fill properly!)
    sender_password = os.getenv("APP_PASS_MAIL")
    if sender_password is None:
        raise EnvironmentError("❗ APP_PASS_MAIL environment variable is not set. "
                               "Please set this environment variable locally or via GitHub Secrets, or on your local environment as an environment variable.")

    leverancier_naam = safe_get(provider_replacements, "LeverancierNaam")
    klant_naam = safe_get(customer_replacements, "KlantNaam")
    klant_job_title = safe_get(customer_replacements, "KlantJobTitle")
    klant_job_reference = safe_get(customer_replacements, "KlantJobReference")
    base_document = f"{base_output_document_path} - {leverancier_naam} - {klant_naam} - {klant_job_title} - {klant_job_reference}"

    email_subject = f"Recht om te vertegenwoordigen documents for '{klant_naam}' for '{klant_job_title}' ({klant_job_reference})"
    email_message = EmailMessage()
    email_message['Subject'] = email_subject
    email_message['From'] = mail_sender_email
    email_message['To'] = email_address
    email_message.set_content(return_html_body(base_document, leverancier_naam, klant_naam, klant_job_title, klant_job_reference), subtype='html')

    for filepath in generated_files:
        abs_full_path = os.path.abspath(filepath)
        if os.path.exists(filepath):
            logging.info(f"ℹ️File at location '{filepath}' found, at absolute path {abs_full_path}")
        else:
            logging.warning(f"⚠️File at location '{filepath}' not found, at absolute path {abs_full_path}!")

        file_basename = os.path.basename(filepath)
        with open(abs_full_path, 'rb') as file:
            file_data = file.read()

            # Choosing the right MIME type explicitly for each file
            if filepath.endswith('.xlsx'):
                maintype, subtype = 'application', 'vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            elif filepath.endswith('.docx'):
                maintype, subtype = 'application', 'vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif filepath.endswith('.pdf'):
                maintype, subtype = 'application', 'pdf'
            else:
                maintype, subtype = 'application', 'octet-stream'

            email_message.add_attachment(
                file_data,
                maintype=maintype,
                subtype=subtype,
                filename=file_basename
            )
    # Send email securely
    try:
        with smtplib.SMTP(mail_smtp_server, mail_smtp_port) as smtp:
            smtp.starttls()
            smtp.login(mail_sender_email, sender_password)
            smtp.send_message(email_message)
        print('✅ Email successfully sent.')
    except Exception as e:
        print('❗ An error occurred:', e)


def set_custom_properties(extracted_dir, provider_replacements, customer_replacements):
    # Now iterate once explicitly since validated
    for key_prov, value_prov in provider_replacements.items():
        set_custom_property(extracted_dir, key_prov, value_prov)
    # Customer replacements (the part you asked about previously)
    for key_cust, value_cust in customer_replacements.items():
        set_custom_property(extracted_dir, key_cust, value_cust)

def set_custom_property(extracted_dir, property_name, property_value):
    custom_props_path = os.path.join(extracted_dir, 'docProps', 'custom.xml')

    tree = ET.parse(custom_props_path)
    root = tree.getroot()

    namespaces = {
        'cp': base_document_namespace_cp,
        'vt': base_document_namespace_vt
    }

    prop = None
    for p in root.findall('cp:property', namespaces):
        if p.get('name') == property_name:
            prop = p
            break

    if prop is not None:
        vt_elem = next(iter(prop))
        vt_elem.text = str(property_value)
    else:
        prop_ids = [int(p.get('pid')) for p in root.findall('cp:property', namespaces)]
        new_pid = max(prop_ids, default=1) + 1

        prop = ET.SubElement(root, '{%s}property' % namespaces['cp'], name=property_name,
                             pid=str(new_pid), fmtid='{D5CDD505-2E9C-101B-9397-08002B2CF9AE}')
        vt_elem = ET.SubElement(prop, '{%s}lpwstr' % namespaces['vt'])
        vt_elem.text = str(property_value)

    tree.write(custom_props_path, encoding='UTF-8', xml_declaration=True, standalone=True)


def extract_docx(docx_path):
    extracted_dir = tempfile.mkdtemp()
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        docx_zip.extractall(extracted_dir)
    return extracted_dir

def repack_docx(extracted_dir, base_document):
    save_as_docx = base_document + ".docx"
    output_docx_path = os.path.abspath(save_as_docx)
    with zipfile.ZipFile(output_docx_path, 'w', zipfile.ZIP_DEFLATED) as docx_zip:
        for foldername, subfolders, filenames in os.walk(extracted_dir):
            for filename in filenames:
                abs_name = os.path.join(foldername, filename)
                arc_name = os.path.relpath(abs_name, extracted_dir)
                docx_zip.write(abs_name, arc_name)
    shutil.rmtree(extracted_dir)


def create_sanitized_replacements(excel_filepath, sheet_name, prefix, optionals=None):
    if optionals :
        sanitized_dict = {}
        sanitized_row_dict = {
            sanitize_spaces_to_variable_name(k): v
            for k, v in optionals.items()
        }
        sanitized_dict[f'{prefix}_0'] = sanitized_row_dict
        return sanitized_dict

    df = pd.read_excel(excel_filepath, sheet_name=sheet_name, header=0)
    if df.empty:
        logging.warning(f"⚠️The provided Excel file '{excel_filepath}' with sheet '{sheet_name}' is empty, No data to process. Please verify the contents.")
        sanitized_dict = {
            sanitize_spaces_to_variable_name(column_name): 'prov_0'
            for column_name in df.columns
        }
    else:
        sanitized_dict = {}
        for index, row in df.iterrows():
            row_dict = row.to_dict()
            # Sanitize KEYS ONLY exactly as you require
            sanitized_row_dict = {
                sanitize_spaces_to_variable_name(k): v
                for k, v in row_dict.items()
            }
            sanitized_dict[f'{prefix}_{index}'] = sanitized_row_dict

    return sanitized_dict


def replace_direct_text(document, provider_replacements, customer_replacements):
    def replace_in_paragraph(paragr, replacem):
        for run in paragr.runs:
            for old, new in replacem.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    logging.info(f"ℹ️'{old}' successfully replaced by '{new}' in target file in paragraphs")

    # Replace in paragraphs
    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, provider_replacements)
        replace_in_paragraph(paragraph, customer_replacements)
    # Replace in tables (cells)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, provider_replacements)
                    replace_in_paragraph(paragraph, customer_replacements)

def replace_images(document):
    replace_images_by_alt_text(document, base_document_image_alt_text_left, image_file_path)
    replace_images_by_alt_text(document, base_document_image_alt_text_right, image_file_path)

def replace_images_by_alt_text(doc, alt_text, new_image_path):
    with open(new_image_path, 'rb') as image_file:
        image_data = image_file.read()
    replaced = False
    for shape in doc.inline_shapes:
        doc_pr = shape._inline.docPr  # protected member accessed due to limitation of python-docx API
        if doc_pr.get('descr') == alt_text:
            rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            doc.part.related_parts[rId]._blob = image_data
            replaced = True
            pass
    if not replaced:
        logging.warning(f"⚠️No image found with alt_text '{alt_text}'")
    else:
        logging.info(f"ℹ️Image with alt_text '{alt_text}' in target file successfully replaced")


def save_to_excel(data_frame, log_data_frame):
    with pd.ExcelWriter(xls_offers_log, mode='w', engine='openpyxl') as log_writer:
        log_data_frame.to_excel(log_writer, sheet_name=xls_offers_log_sheetname, index=False)
    with pd.ExcelWriter(xls_offers_customer, mode='w', engine='openpyxl') as cust_writer:
        data_frame.to_excel(cust_writer, sheet_name=xls_offers_customer_sheetname, index=False)
    return log_data_frame

def safely_read_excel(excel_file, sheet_name, description):
    """
    Intelligent wrapper function clearly handling Excel read and logging errors dynamically.
    """
    try:
        return pd.read_excel(excel_file, sheet_name)
    except ValueError as e:
        error_msg = f"❌ValueError occurred reading {description} Excel file '{excel_file}', sheet '{sheet_name}': {e}"
        logging.error(error_msg)
        raise
    except Exception as e:
        error_msg = f"❌Unexpected error reading {description} Excel file '{excel_file}', sheet '{sheet_name}': {e}"
        logging.error(error_msg)
        raise


def convert_to_pdf(base_document):
    convert_from_docx = base_document + ".docx"
    save_as_pdf = base_document + ".pdf"

    if not os.path.exists(output_path):
        os.makedirs(output_path)
    abs_full_path_convert_from_docx = os.path.abspath(convert_from_docx)
    abs_full_path_save_as_pdf = os.path.abspath(save_as_pdf)
    abs_output_path = os.path.abspath(output_path)

    try:
        subprocess.run([
            'soffice',
            '--headless',
            '--convert-to',
            'pdf',
            '--outdir',
            output_path,
            convert_from_docx
        ], check=True)
        logging.info("ℹ️Word file: " + convert_from_docx + " with absolute path: " + abs_full_path_convert_from_docx)
        logging.info("ℹ️Successfully converted to Pdf file: " + save_as_pdf + " with absolute path: " + abs_full_path_save_as_pdf)
        files = os.listdir(abs_output_path)
        logging.info(f"📂 Explicitly listing files from '{abs_output_path}':")
        if files:
            for file in files:
                logging.info(f"    - {file}")
        else:
            logging.warning(f"📭 Directory '{abs_output_path}' explicitly exists but is empty!")
    except FileNotFoundError:
        logging.error(f"🚨 Directory '{abs_output_path}' not found!")
    except subprocess.CalledProcessError as e:
        logging.error(f"❌Could not convert PDF: {e}")

def convert_to_pdf_traditional(base_document):
    save_as_docx = base_document + ".docx"
    save_as_pdf = base_document + ".pdf"
    try:
        # Convert the output
        convert(save_as_docx, save_as_pdf)
        logging.info("ℹ️Word file: " + save_as_docx)
        logging.info("ℹ️Successfully converted to Pdf file: " + save_as_pdf)
    except Exception as e:
        logging.error(f"❌Failed to convert {save_as_docx} to PDF: {e}")
        if os.path.exists(save_as_docx):  # to avoid stale files
            logging.info(f"ℹ️Cleaning up incomplete conversion file {save_as_docx}")
            os.remove(save_as_docx)
        raise


def open_document(base_document):
    save_as_docx = base_document + ".docx"
    return Document(save_as_docx)

def save_document(base_document, document):
    save_as_docx = base_document + ".docx"
    document.save(save_as_docx)


def sanitize_filename(filename_part):
    return re.sub(r'[*<>:"/\\|?]', '_', filename_part).strip()

def sanitize_spaces_to_variable_name(any_string):
    return any_string.replace(" ", "")

def safe_get(row, column_name, default='unknown'):
    return sanitize_filename(str(row[column_name]).strip()) if column_name in row else default

def build_base_document_to_save(provider_replacements, customer_replacements):
    leverancier_naam = safe_get(provider_replacements, "LeverancierNaam")
    klant_naam = safe_get(customer_replacements, "KlantNaam")
    klant_job_title = safe_get(customer_replacements, "KlantJobTitle")
    klant_job_reference = safe_get(customer_replacements, "KlantJobReference")
    base_document = f"{base_output_document_path} - {leverancier_naam} - {klant_naam} - {klant_job_title} - {klant_job_reference}"
    return base_document

def return_html_body(base_document, leverancier_naam, klant_naam, klant_job_title, klant_job_reference):
    return f"""
    <body>
        <h2>Recht om te vertegenwoordigen</h2>
        <table>
            <tr>
                <td><b>Leverancier Naam</b></td>
                <td>{leverancier_naam}</td>
            </tr>
            <tr>
                <td><b>Klant Naam</b></td>
                <td>{klant_naam}</td>
            </tr>
            <tr>
                <td><b>Klant JobTitle</b></td>
                <td>{klant_job_title}</td>
            </tr>
            <tr>
                <td><b>Klant JobReference</b></td>
                <td>{klant_job_reference}</td>
            </tr>
        </table>

        <h3>Documents Attached:</h3>
        <table>
            <tr>
                <td><b>MSWord</b></td>
                <td>{base_document}.docx</td>
            </tr>
            <tr>
                <td><b>Pdf</b></td>
                <td>{base_document}.pdf</td>
            </tr>
        </table>
    </body>
    """

if __name__ == '__main__':
    # Parse command-line arguments clearly here
    parser = argparse.ArgumentParser(description="Generate Offer Documents with optional verbose logging, and possibly customer arguments.")
    parser.add_argument("--klantNaam", help="Klant Naam")
    parser.add_argument("--klantJobTitle", help="Klant JobTitle")
    parser.add_argument("--klantJobReference", help="Klant JobReference")
    parser.add_argument("--emailRecipient", help="Email Recipient")
    parser.add_argument("--uploadDropbox", help="Upload Dropbox")
    parser.add_argument("--verbose", action="store_true", help="Enable verbose output")

    args = parser.parse_args()

    if args.klantNaam and args.klantJobTitle and args.klantJobReference:
        optional_args = {
            'KlantNaam': args.klantNaam,
            'KlantJobTitle': args.klantJobTitle,
            'KlantJobReference': args.klantJobReference,
            'EmailRecipient': args.emailRecipient,
            'UploadDropbox': args.uploadDropbox
        }
        _main(
            verbose=args.verbose,
            optional_args=optional_args
        )
    else:
        _main(
            verbose=args.verbose
        )
